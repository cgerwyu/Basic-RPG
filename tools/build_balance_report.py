from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Basic_RPG_Classes_PvP_PvE_Rebalance_RU.docx"

PAGE_WIDTH_DXA = 12240
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
HEADER_FILL = "E8EEF5"
LIGHT_FILL = "F4F6F9"
PALE_BLUE = "EDF4FA"
PALE_GOLD = "FFF8E8"
PALE_RED = "FFF1F1"
NAVY = RGBColor(31, 77, 120)
BLUE = RGBColor(46, 116, 181)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(92, 103, 114)
GOLD = RGBColor(122, 90, 0)
RED = RGBColor(155, 28, 28)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=TABLE_INDENT_DXA):
    assert sum(widths) == CONTENT_WIDTH_DXA, widths
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = borders.find(qn(f"w:{side}"))
        if edge is None:
            edge = OxmlElement(f"w:{side}")
            borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), "4")
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), "C8D1DC")

    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for side, value in (("top", 80), ("start", 120), ("bottom", 80), ("end", 120)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            cell.width = Inches(widths[idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_run(run, size=11, color=INK, bold=False, italic=False, font="Calibri"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def style_para(p, after=6, before=0, line=1.25, align=None):
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if align is not None:
        p.alignment = align
    return p


def add_para(doc, text="", *, bold=False, italic=False, size=11, color=INK,
             after=6, before=0, line=1.25, align=None, style=None):
    p = doc.add_paragraph(style=style)
    style_para(p, after=after, before=before, line=line, align=align)
    if text:
        set_run(p.add_run(text), size=size, color=color, bold=bold, italic=italic)
    return p


def add_rich_para(doc, parts, *, after=6, before=0, line=1.25, align=None):
    p = doc.add_paragraph()
    style_para(p, after=after, before=before, line=line, align=align)
    for text, options in parts:
        set_run(p.add_run(text), **options)
    return p


def add_hyperlink(paragraph, text, url):
    rel_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2E74B5")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([r_fonts, color, underline])
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run(run, size=9, color=MUTED)


def add_bottom_border(paragraph, color="D7DEE8", size="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def add_numbering(doc, num_fmt, marker):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(n.get(qn("w:abstractNumId"))) for n in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), marker)
    jc = OxmlElement("w:lvlJc")
    jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend([tabs, ind, spacing])
    lvl.extend([start, fmt, lvl_text, jc, p_pr])
    abstract.append(lvl)
    # OOXML requires every abstractNum before the first concrete num. Appending
    # a new abstract after the built-in num elements makes Word silently repair
    # the file and can turn bullets into decimal lists during PDF export.
    first_num = numbering.find(qn("w:num"))
    if first_num is None:
        numbering.append(abstract)
    else:
        numbering.insert(numbering.index(first_num), abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_list_item(doc, text, num_id, *, bold_prefix=None):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_node])
    p_pr.append(num_pr)
    style_para(p, after=4, line=1.25)
    if bold_prefix and text.startswith(bold_prefix):
        set_run(p.add_run(bold_prefix), bold=True)
        set_run(p.add_run(text[len(bold_prefix):]))
    else:
        set_run(p.add_run(text))
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.page_break_before = False
    set_run(p.add_run(text), size={1: 16, 2: 13, 3: 12}[level],
            color=BLUE if level < 3 else NAVY, bold=True)
    return p


def add_table(doc, headers, rows, widths, *, font_size=9.1, header_fill=HEADER_FILL):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    repeat_header(table.rows[0])
    prevent_row_split(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        style_para(p, after=0, line=1.05)
        set_run(p.add_run(header), size=font_size, color=INK, bold=True)
    for row in rows:
        table_row = table.add_row()
        prevent_row_split(table_row)
        cells = table_row.cells
        for idx, value in enumerate(row):
            set_cell_width(cells[idx], widths[idx])
            p = cells[idx].paragraphs[0]
            style_para(p, after=0, line=1.08)
            set_run(p.add_run(str(value)), size=font_size, color=INK)
    add_para(doc, "", after=2, line=1.0)
    return table


def add_table_source(doc, text, links):
    p = doc.add_paragraph()
    style_para(p, before=4, after=4, line=1.05)
    set_run(p.add_run(text), size=8.5, color=MUTED, italic=True)
    for idx, (label, url) in enumerate(links):
        if idx > 0:
            set_run(p.add_run(" · "), size=8.5, color=MUTED)
        add_hyperlink(p, label, url)
    return p


def add_callout(doc, title, body, fill=LIGHT_FILL, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    # A one-cell callout behaves as a labelled note. Marking its single row as
    # a header gives assistive readers a stable semantic entry point and also
    # makes the row repeat safely if Word ever splits the note across pages.
    repeat_header(table.rows[0])
    prevent_row_split(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    style_para(p, after=3, line=1.15)
    set_run(p.add_run(title), size=11, color=accent, bold=True)
    p2 = cell.add_paragraph()
    style_para(p2, after=0, line=1.18)
    set_run(p2.add_run(body), size=10.2, color=INK)
    add_para(doc, "", after=3, line=1.0)
    return table


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for level, size, color, before, after in (
        (1, 16, BLUE, 18, 10),
        (2, 13, BLUE, 14, 7),
        (3, 12, NAVY, 10, 5),
    ):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    hp = header.paragraphs[0]
    style_para(hp, after=0, line=1.0)
    set_run(hp.add_run("BASIC RPG CLASSES  ·  БАЛАНС V2"), size=8.5, color=MUTED, bold=True)
    add_bottom_border(hp)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_para(fp, before=0, after=0, line=1.0)
    set_run(fp.add_run("Страница "), size=9, color=MUTED)
    add_page_field(fp)


def section_transition(doc, *, force_page=False):
    if force_page:
        doc.add_page_break()
        return
    # Compact-reference layout: use the remaining page space. Heading
    # keep-with-next rules still prevent an orphaned section title.
    add_para(doc, "", after=6, line=1.0)


def build():
    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0])
    doc.core_properties.title = "Basic RPG Classes — PvP/PvE баланс v2"
    doc.core_properties.subject = "Исследовательский дизайн-отчёт по классам, режимам и прогрессии"
    doc.core_properties.author = "OpenAI Codex"
    doc.core_properties.keywords = "Minecraft, NeoForge, PvP, PvE, class balance, game design"

    bullet_id = add_numbering(doc, "bullet", "•")
    number_id = add_numbering(doc, "decimal", "%1.")

    # Editorial cover: named override over compact_reference_guide.
    add_para(doc, "", after=72, line=1.0)
    add_para(doc, "ИССЛЕДОВАТЕЛЬСКИЙ ДИЗАЙН-ОТЧЁТ", bold=True, size=10.5, color=GOLD,
             after=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Basic RPG Classes", bold=True, size=30, color=INK,
             after=8, align=WD_ALIGN_PARAGRAPH.CENTER, line=1.0)
    add_para(doc, "PvP/PvE баланс v2", bold=True, size=16, color=NAVY,
             after=4, align=WD_ALIGN_PARAGRAPH.CENTER, line=1.0)
    add_para(doc, "Личный навык важнее одной сильной кнопки", italic=True, size=12, color=MUTED,
             after=52, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_callout(
        doc,
        "Главное решение",
        "Не пытаться балансировать весь мод одним множителем. Один и тот же навык сохраняет роль, но получает отдельные числовые профили для обычного PvE, боссов, дуэлей, массовых войн и open-world PvP.",
        fill=PALE_BLUE,
        accent=BLUE,
    )
    add_para(doc, "25 августа 2026", bold=True, size=11, color=NAVY,
             after=3, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Основано на аудите текущего кода и первичных источниках Valve, ArenaNet, Blizzard, Riot и NeoForge",
             size=9.3, color=MUTED, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)

    section_transition(doc, force_page=True)
    add_heading(doc, "1. Итог для принятия решения", 1)
    add_para(doc, "Текущая проблема — не просто завышенные цифры. У Охотника и Мага накопились несколько независимых преимуществ, которые складываются: дальность, мобильность, burst, контроль и безопасный выход. Воин получает силу только после контакта, но не имеет достаточно надёжного и ограниченного по риску способа этот контакт создать.")
    add_callout(
        doc,
        "Вердикт",
        "Охотнику не нужно удалять классовую мобильность, а Воину не нужен свободный телепорт. Нужно связать мобильность и усиленные выстрелы Охотника одной Выносливостью, дать Воину точный intercept и короткое anti-mobility окно, а Магу — высокий урон только через телеграфируемые и прерываемые заклинания.",
        fill=PALE_GOLD,
        accent=GOLD,
    )
    add_table(
        doc,
        ["Проблема", "Почему ломает бой", "Решение v2"],
        [
            ("Один PvP-множитель", "Не ограничивает очередь, DoT и канал", "Per-cast/per-target budget + режимные профили"),
            ("Hunter: mobility + burst", "Побег не конкурирует с атакой", "Одна Выносливость для Dash, Windrun, stealth и усиленных стрел"),
            ("Warrior не догоняет", "Сила существует только после контакта", "Intercept, ranged guard, короткий mobility lock"),
            ("Mage делает всё", "AoE, burst, heal, cleanse, shield, blink", "Heal → Priest; cast risk; один Blink"),
            ("Veteran против новичка", "Уровень/gear решают до начала боя", "Arena normalization; GvG compression; opt-in и anti-camp"),
        ],
        [2100, 3200, 4060],
        font_size=9.2,
    )

    add_heading(doc, "1.1. Что уже исправлено в коде", 2)
    for item in (
        "Четыре стрелковых навыка Охотника теперь можно подготовить без классового лука именно в основной руке; эффект применяется следующим выстрелом классового лука.",
        "Dash больше не отвечает общей ошибкой при вертикальном взгляде и всегда использует устойчивое горизонтальное направление.",
        "Одна очередь Multishot поражает каждую цель не более одного раза; один Arrow Rain поражает каждого игрока не более одного раза.",
        "Hunter Mana переименована в Выносливость; её пул и пассивы сжаты, а восстановление настроено под общий расход offense/mobility.",
        "Magic Shield теперь даёт 1.2 секунды полной неуязвимости вместо пятисекундной защиты и очищает не более пяти союзников.",
        "Meteor предупреждается 2.5 секунды; Meteor и Sky Rays прерываются движением или полученным уроном.",
        "Fireball больше не выпускает три полноценных снаряда за одну цену; чрезмерные пассивы здоровья, маны, натяжения и урона стрел снижены.",
        "Проект собран без ошибок; актуальный JAR создаётся в build/libs.",
    ):
        add_list_item(doc, item, bullet_id)

    add_heading(doc, "1.2. Что сознательно не сделано одним hotfix", 2)
    add_para(doc, "Новые классы, manager дуэлей/гильдий, бижутерия, boss break bar и data-driven таблицы требуют миграции данных и отдельного тестового цикла. Их опасно смешивать с исправлением неработающих навыков: ошибка миграции может повредить существующие сохранения.")

    section_transition(doc)
    add_heading(doc, "2. Что подтверждает исследование", 1)
    add_table(
        doc,
        ["Проверенный принцип", "Наблюдение в первичном источнике", "Применение в моде"],
        [
            ("Разные режимы — разные числа", "GW2 разделяет PvE, PvP и WvW balance", "Пять CombatProfile поверх одного skill identity"),
            ("Дальник слабеет вблизи", "Drow теряет часть пассивной силы рядом с врагом", "Hunter получает close-range pressure penalty"),
            ("Милишник ловит ограниченно", "Sven/Ursa: catch → опасное melee-окно", "Warrior Intercept без свободного teleport"),
            ("Большая магия предупреждается", "Invoker AoE имеет delay/setup/cost", "Meteor/channeled ult можно сорвать или покинуть"),
            ("Контроль босса координируется", "ArenaNet Defiance bar регенерирует и видима", "Poise damage вместо stun immunity"),
            ("Формальный PvP равняет старт", "GW2 нормализует уровень/gear", "Дуэльный tournament template"),
            ("Ресурс создаёт цену промаха", "Riot: mana задаёт ebb/flow; второй bar нужен не всегда", "Hunter начинает с одной общей Выносливости"),
        ],
        [2250, 3520, 3590],
        font_size=8.8,
    )
    add_table_source(doc, "Основные источники: ", [
        ("ArenaNet о разделении режимов", "https://www.guildwars2.com/en/news/updates-from-the-systems-team/"),
        ("Dota 2 heroes", "https://www.dota2.com/heroes"),
        ("Riot о ресурсах", "https://www.leagueoflegends.com/en-us/news/dev/ask-riot-manaless-champions/"),
    ])

    add_heading(doc, "2.1. Важная поправка к аналогии с Dota 2", 2)
    add_para(doc, "Dota 2 не доказывает, что Meteor обязан держать персонажа неподвижным ровно три секунды. Она доказывает необходимость телеграфа, цены и setup. Прерываемый 2.5-секундный cast — отдельное решение под геометрию Minecraft, где нет тумана войны и предметов мобильности Dota.")
    add_para(doc, "Точно так же Dota не требует делать Priest почти без урона: Oracle и Dazzle способны драться. Для мода правильнее дать Priest низкий sustained damage, достаточный для solo PvE, но не burst, сравнимый с Mage/Hunter.")

    section_transition(doc)
    add_heading(doc, "3. Архитектура режимов", 1)
    add_para(doc, "Один SkillId должен сохранять название, назначение, анимацию и основные контрходы. Меняются коэффициенты, длительности контроля, target cap, ресурс и допустимые типы целей.")
    add_table(
        doc,
        ["Профиль", "Цель", "Экипировка", "Ключевые правила"],
        [
            ("PVE_NORMAL", "Solo/party против обычных мобов", "Полная сила прогрессии", "Исследовательская мобильность; широкий AoE"),
            ("PVE_BOSS", "Vanilla и будущие боссы", "Полная сила прогрессии", "CC → poise; boss heal/lifesteal multipliers"),
            ("PVP_DUEL", "1v1 и малая арена", "Standard template, ±5%", "Самые строгие burst/sustain budgets"),
            ("PVP_WAR", "Гильдии, 5v5–20v20", "35% бонуса сверх baseline", "AoE falloff; support cap 5; anti-stack"),
            ("PVP_OPEN_WORLD", "Добровольное PvP в мире", "Сжатая earned gear", "Opt-in; sanctuary; novice/respawn; anti-camp"),
        ],
        [1700, 2400, 2100, 3160],
        font_size=8.8,
    )

    add_heading(doc, "3.1. Целевые бюджеты alpha", 2)
    add_table(
        doc,
        ["Показатель", "Дуэль", "Массовый PvP", "PvE/босс"],
        [
            ("Median TTK", "14–22 с", "Focus 3+ игроков: 4–8 с", "Задаётся encounter tier"),
            ("Instant non-ultimate", "≤15% effective HP", "≤12% на цель", "По типу моба"),
            ("Telegraphed skillshot", "≤22%", "≤18% на цель", "До полного PvE значения"),
            ("Полностью попавшая ultimate", "≤35%", "AoE falloff + cap", "Высокий урон за механику"),
            ("Один cast за 2 с", "Hard cap 40%", "Hard cap 35%/цель", "Boss profile"),
            ("Hard CC", "0.5–0.9 с + DR", "0.4–0.7 с + DR", "Poise damage"),
            ("Support targets", "1–2", "5", "Party size/profile"),
        ],
        [2550, 2160, 2420, 2230],
        font_size=8.9,
    )
    add_callout(doc, "Почему cap считается на cast", "Семь стрел, пять импульсов и три fireball — это одна кнопка. Ограничение каждого отдельного damage event позволяет безнаказанно обойти общий cap количеством событий.", fill=PALE_RED, accent=RED)

    add_heading(doc, "3.2. Контроль и иммунитет", 2)
    for item in (
        "Hard-control одной категории: 100% длительности → 50% → иммунитет; reset после 8 секунд без нового контроля.",
        "Mobility lock считается soft-control: длится 1.25–1.75 секунды, не складывается и имеет ясную иконку/частицы.",
        "Неуязвимость — только короткая reaction-кнопка. Групповая универсальная неуязвимость запрещена; Priest защищает лишь от конкретного типа угрозы.",
        "Boss не получает обычный stun до слома видимой poise bar; величина шкалы масштабируется по tier и числу игроков.",
    ):
        add_list_item(doc, item, bullet_id)

    section_transition(doc)
    add_heading(doc, "4. Ресурсы, ранги и экипировка", 1)
    add_table(
        doc,
        ["Класс", "Ресурс", "Что ограничивает", "Регенерация и конфликт"],
        [
            ("Warrior", "Ярость 0–100", "Defense, catch, AoE payoff", "Получает за контакт/блок; вне боя постепенно теряет"),
            ("Mage", "Мана 150 baseline", "Poke, control, burst, blink", "Combat regen низкий; мощный cast оставляет без offense"),
            ("Hunter", "Выносливость 100", "Dash, Windrun, stealth, shot modifiers", "Побег напрямую отнимает ресурс у следующей атаки"),
            ("Priest", "Мана/Вера", "Healing throughput и dispel", "Overheal не возвращает ресурс; длинный бой истощает"),
            ("Summoner", "Command/Essence", "Resummon и активные приказы", "Потеря summon создаёт реальный downtime"),
        ],
        [1500, 1700, 2740, 3420],
        font_size=8.9,
    )
    add_callout(doc, "Почему Hunter не получает две независимые полосы сразу", "Mana + Stamina звучат глубже, но позволяют тратить один пул на идеальный burst и сохранять второй на гарантированный побег. На первом этапе одна Выносливость создаёт более сильный выбор. Второй ресурс допустим только после телеметрии, если он будет конфликтовать, а не дублировать.", fill=PALE_GOLD, accent=GOLD)

    add_heading(doc, "4.1. Где должна находиться сила прогрессии", 2)
    add_table(
        doc,
        ["Источник", "PvE", "Дуэль", "GvG/open world"],
        [
            ("Skill rank 1→15", "+20–25% effect; −10–15% cost/CD", "Турнирный effective rank", "Сжатый earned rank"),
            ("Оружие", "Главный offensive growth", "Стандартный профиль", "35% bonus сверх baseline; cap"),
            ("Броня", "Главный defensive growth", "Стандартный профиль", "35% bonus; defense cap выше offense"),
            ("Бижутерия", "Mana/regen/utility tradeoffs", "Выбор из равных PvP-вариантов", "Эффекты strongest-only, без stack"),
            ("Пассивы", "Небольшая идентичность", "Сохраняются умеренно", "Не более 10–20% baseline"),
        ],
        [1800, 2830, 2100, 2630],
        font_size=8.8,
    )

    section_transition(doc)
    add_heading(doc, "5. Warrior — антикит без телепорта", 1)
    add_para(doc, "Воин должен оставаться самым медленным в свободном перемещении, но самым опасным после честно выигранного сближения. Его мастерство — выбрать момент входа, не потратить защиту впустую и удержать контакт.")
    add_table(
        doc,
        ["Навык", "Цель v2", "Дуэль", "GvG / PvE boss", "Контригра"],
        [
            ("Whirlwind", "AoE payoff после входа", "5–7 HP raw; lifesteal cap", "До 5 PvP целей; высокий PvE clear", "Выйти из melee/radius"),
            ("Fortify → Iron Advance", "Дойти под огнём", "4 с: −40% ranged, slow/KB resist", "Self-only defense; party utility отдельно", "Переждать окно; атаковать melee"),
            ("Provoke/Challenge", "Защитить союзника", "−25–28% урона по другим, reveal", "Cap 5; PvE taunt", "Бить Warrior или disengage"),
            ("Ground Stun", "Поймать ошибку позиции", "0.6–0.8 с hard CC + 1.5 с lock", "Poise damage boss", "Телеграфируемый cone"),
            ("Shield Bash/Intercept", "Точный gap closer", "Узкий target/LoS; 1.5 с mobility lock", "Прыжок к priority target", "Sidestep/obstacle; промах наказуем"),
            ("Battle Cry", "Короткий team window", "+10% offense/tenacity 5 с", "Cap 5; strongest-only", "Disengage или counter-ultimate"),
        ],
        [1550, 2050, 2100, 2010, 1650],
        font_size=8.3,
    )
    add_heading(doc, "5.1. Необходимые ограничения", 2)
    for item in (
        "Вампиризм работает от AoE, но PvP-heal имеет per-tick/per-cast cap; толпа не делает Воина математически бессмертным.",
        "Intercept не ищет цель за стеной и не корректирует промах после запуска.",
        "Anti-mobility не запрещает обычное W/A/S/D и прыжок; оно блокирует только Blink/Dash/Windrun/Climb/Glide.",
        "Воин не получает постоянную пассивную защиту от дальнего урона — это активное окно с КД 24–30 секунд.",
    ):
        add_list_item(doc, item, bullet_id)

    section_transition(doc)
    add_heading(doc, "6. Mage — риск ради AoE и контроля", 1)
    add_para(doc, "Маг остаётся лучшим в массовой магии, но больше не совмещает top burst, полноценный heal, cleanse, длительную защиту и два бесплатных Blink. Большая сила существует только через читаемый cast и расход маны.")
    add_table(
        doc,
        ["Навык", "Цель v2", "Дуэль", "GvG / PvE", "Контригра"],
        [
            ("Fireball", "Основной skillshot", "1 projectile; 12–18% HP", "Небольшой splash; PvE coefficient", "Dodge/block; промах стоит mana"),
            ("Heal (legacy)", "Миграция в Priest", "До миграции — слабый self/party heal", "После миграции отсутствует у Mage", "Pressure и anti-heal"),
            ("Blink", "Один reposition", "1 charge; 9–11 с; attack lockout", "Mana cost; no chained escape", "Mobility lock; predict endpoint"),
            ("Magic Shield", "Reaction + cleanse", "0.75–1.2 с invulnerability", "Self invuln + cleanse до 5 allies", "Bait; затем длинный КД"),
            ("Frost Nova", "Self-peel", "Малый damage; slow/root DR", "Zone control; boss poise", "Не входить одновременно всей группой"),
            ("Chain Lightning", "Multi-target pressure", "Умеренный single-target", "Falloff по jumps; target cap", "Spread; LoS"),
            ("Sky Rays", "Прерываемый channel", "3 с, fixed total budget", "Сектор/LoS; PvE targets выше", "Damage, movement, stun"),
            ("Meteor", "Telegraphed ultimate", "2.5 с cast; ≤35% HP", "AoE falloff; 40–50% base mana", "Выйти, прервать, force move"),
        ],
        [1450, 1900, 2100, 2160, 1750],
        font_size=8.2,
    )
    add_callout(doc, "Magic Shield — не прежняя пятисекундная неуязвимость", "Короткое окно позволяет переиграть точный burst, но не даёт Магу продолжать безнаказанную ротацию. Группе передаётся cleanse, а не универсальная invulnerability.", fill=PALE_BLUE, accent=BLUE)

    section_transition(doc)
    add_heading(doc, "7. Hunter — мобильность как ограниченный бюджет", 1)
    add_para(doc, "Hunter остаётся самым мобильным классом, но больше не может одновременно использовать Dash, Speed, Climb и Camouflage, а затем сохранить полный offensive burst. Дистанция — награда за управление Выносливостью, а не постоянное состояние.")
    add_table(
        doc,
        ["Навык", "Цель v2", "Дуэль", "GvG / PvE", "Контригра"],
        [
            ("Dash", "Короткий reposition", "1 charge; 8–10 с; 24 stamina", "Вне PvP возможны talent bonuses", "Предсказать; mobility lock"),
            ("Windrun", "Окно kite", "Speed I 3–4.5 с; −20% direct damage", "PvE direct mob evade", "Переждать; AoE/DoT"),
            ("Camouflage", "Смена позиции", "3–5 с; break on damage/action", "Сброс mob aggro", "Reveal, proximity, AoE"),
            ("Multishot", "Clear, не shotgun", "Одна цель — одно попадание/volley", "3/5/7 targets/arrows", "Не стоять группой"),
            ("Frost Arrows", "Цена за контроль", "Stamina/shot; slow DR", "Boss poise/soft CC coefficient", "Force resource starvation"),
            ("Arrow Rain", "Area denial", "Один damage hit/player/cast", "Больше PvE hits; data-driven boss cap", "Покинуть телеграф"),
            ("Power Shot", "Высокий skillshot", "1–1.5 с charge; ≤22% HP", "Pierce/falloff; boss damage", "Interrupt, sidestep, cover"),
            ("Climbing", "Исследование мира", "Off на arena и в PvP combat", "Полностью доступен в PvE", "Начать бой до стены; lock"),
        ],
        [1450, 1900, 2110, 2150, 1750],
        font_size=8.15,
    )
    add_heading(doc, "7.1. Штраф вблизи", 2)
    add_para(doc, "Когда вражеский игрок находится примерно в 3–4 блоках, Hunter теряет часть пассивного бонуса натяжения/точности. Это не выключает лук полностью, но делает выигранное Warrior сближение значимым. Camouflage и Dash в этот момент остаются возможными только при наличии Выносливости и отсутствии mobility lock.")

    section_transition(doc)
    add_heading(doc, "8. Priest — следующий класс", 1)
    add_para(doc, "Priest нужен прежде Summoner: перенос Heal снимает перегрузку Mage и одновременно создаёт востребованную party-role для PvE/GvG. Однако класс не должен быть беспомощным в solo PvE или бессмертным в 1v1.")
    add_table(
        doc,
        ["Навык", "Функция", "PvP-ограничение", "PvE-ценность / контригра"],
        [
            ("Smite", "Низкий sustained damage", "Нет burst-combo", "Достаточно для обычных solo mobs"),
            ("Mend", "Cast heal одной цели", "Cast interrupt; combat coefficient", "Основной tank heal"),
            ("Renew", "HoT", "Не stack; dispellable", "Сглаживает boss pressure"),
            ("Purify", "Снять debuff", "1 category/target; meaningful CD", "Снимает mechanics-tagged effects"),
            ("Barrier", "Типовая защита", "Physical ИЛИ magic, не всё", "Подготовка к telegraphed boss hit"),
            ("Sanctuary", "Healing zone/object", "Уничтожаемый/покидаемый; cap 5", "Сильная позиционная поддержка"),
            ("Resurrection", "Вернуть союзника", "Запрещено в дуэли; долгий war channel", "Доступно по encounter rules"),
        ],
        [1600, 2400, 2510, 2850],
        font_size=8.6,
    )
    add_callout(doc, "Роль Priest в 1v1", "Он выигрывает не burst, а экономикой маны, interrupt и правильным dispel. Если healing throughput позволяет бесконечно обнулять входящий урон, класс сломан даже при низком собственном DPS.", fill=PALE_GOLD, accent=GOLD)

    add_heading(doc, "8.1. Правила массового лечения", 2)
    for item in (
        "Приоритет целей: self → party/subgroup → guild squad → остальные союзники по расстоянию.",
        "Competitive cap — 5 целей. Несколько Priest не складывают один и тот же Barrier/HoT; работает strongest-only или refresh.",
        "Overheal не превращается автоматически в shield. PvP combat вводит healing coefficient/dampening.",
        "Сильное спасение имеет уничтожаемый объект, cast time, узкую область или большой cooldown — минимум один явный контрход.",
    ):
        add_list_item(doc, item, bullet_id)

    section_transition(doc)
    add_heading(doc, "9. Summoner — после Priest и boss API", 1)
    add_para(doc, "Summoner особенно опасен для массового PvP и производительности сервера. Поэтому его дизайн начинается не со списка десяти существ, а с жёсткого лимита, поводка и понятной цены потери ключевого призыва.")
    add_table(
        doc,
        ["Правило", "Рекомендуемый baseline", "Зачем"],
        [
            ("Основной summon", "Ровно 1", "Читаемость целей и нагрузка AI"),
            ("Временный summon", "Не более 1 одновременно", "Не создавать армию в GvG"),
            ("Leash", "24–32 блока; вне leash не атакует", "Владелец обязан рисковать позицией"),
            ("PvP damage", "50–65% PvE coefficient", "Pet не выигрывает бой без команд"),
            ("AoE resistance", "Частичная, не immunity", "Не умирать случайно, но иметь counterplay"),
            ("Смерть summon", "Backlash или resummon 45–90 с", "Потеря является настоящей ошибкой"),
            ("Boss control", "Запрещён по умолчанию", "Не ломать будущие encounter mechanics"),
        ],
        [2100, 2700, 4560],
        font_size=8.9,
    )

    section_transition(doc)
    add_heading(doc, "10. Дуэли, войны и защита новичков", 1)
    add_heading(doc, "10.1. Дуэль / арена", 2)
    for item in (
        "Добровольное подтверждение обоих игроков; snapshot инвентаря и состояний; безопасное восстановление после матча.",
        "Стандартный class template: фиксированная броня/оружие, равный budget очков, одинаковые consumables.",
        "Отдельная arena map без climbing exploits; сброс ресурсов/cooldowns перед раундом.",
        "Рейтинг учитывает только нормализованные матчи; open-world kill не влияет на duel MMR.",
    ):
        add_list_item(doc, item, bullet_id)

    add_heading(doc, "10.2. Guild War / массовый PvP", 2)
    for item in (
        "Earned gear сохраняется, но bonus сверх baseline умножается примерно на 0.35 и имеет offense/defense caps.",
        "AoE damage использует target cap/falloff; support cap равен 5 и приоритизирует subgroup.",
        "Friendly fire и союзники определяются нативным war manager, а не только scoreboard prefix.",
        "Reward даётся за objectives, assist, healing/CC contribution, а не только final blow — иначе Hunter burst доминирует экономику.",
    ):
        add_list_item(doc, item, bullet_id)

    add_heading(doc, "10.3. Open-world и anti-gank", 2)
    add_table(
        doc,
        ["Механика", "Правило"],
        [
            ("Opt-in War Mode", "PvP включается добровольно в safe hub; выключается вне combat"),
            ("Sanctuary", "Spawn, class hub, starter villages и tutorial zones не допускают PvP"),
            ("Novice protection", "Учитывает class level и effective gear score; агрессия снимает защиту"),
            ("Respawn/login", "20 с / 8 с; атака игрока снимает щит немедленно"),
            ("Anti-camp", "Повторное убийство одной цели не даёт награду; серийный убийца отмечается bounty"),
            ("Power gap", "Сильный разрыв → запрет первого удара либо жёсткая compression, не скрытый +500% buff"),
        ],
        [2300, 7060],
        font_size=9.0,
    )

    section_transition(doc)
    add_heading(doc, "11. PvE, боссы и будущий мод существ", 1)
    add_para(doc, "Классы не должны знать Java-класс каждого будущего моба. NeoForge позволяет группировать EntityType тегами и прикреплять reloadable data map. Значит, новый мод существ сможет объявить категорию и коэффициенты без патча основного мода.")
    add_table(
        doc,
        ["Data-driven поле", "Пример", "Использование"],
        [
            ("category tag", "normal / elite / boss / summon", "Выбор профиля навыка"),
            ("poise_max", "200 solo; scaling per player", "Break bar и hard CC"),
            ("control_multiplier", "0.5 elite; 0 boss direct", "CC → poise damage"),
            ("lifesteal_multiplier", "1.0 mob; 0.35 boss", "Не давать бесконечный sustain"),
            ("aoe_damage_multiplier", "По encounter tier", "Mass-clear и boss damage отдельно"),
            ("summon_policy", "forbidden / cosmetic / controllable", "Безопасность Summoner"),
            ("heal_received_multiplier", "Undead/construct rules", "Priest и будущие механики"),
        ],
        [2600, 3060, 3700],
        font_size=8.8,
    )
    add_table_source(doc, "Техническая основа: ", [
        ("NeoForge Tags", "https://docs.neoforged.net/docs/resources/server/tags/"),
        ("NeoForge Data Maps", "https://docs.neoforged.net/docs/1.21.5/resources/server/datamaps/"),
    ])

    add_heading(doc, "11.1. Boss break bar", 2)
    for item in (
        "Синяя/бирюзовая шкала видима под HP босса; skill tooltip показывает poise damage.",
        "Hard CC не оглушает босса напрямую, а снимает poise. После слома открывается 2–5-секундное окно stagger/exposed.",
        "Шкала регенерирует, если группа тратит контроль несогласованно; strength масштабируется после заданного числа игроков.",
        "Особые атаки босса могут временно открывать break bar: успешный break прерывает смертельную механику.",
    ):
        add_list_item(doc, item, bullet_id)

    add_heading(doc, "11.2. Сохранение solo и party PvE", 2)
    add_para(doc, "Warrior и Priest не обязаны сравниваться с Mage/Hunter только по DPS. Для обычного solo-контента им нужен приемлемый TTK; для boss-party их меньший урон окупается tanking, break, cleanse и healing. В таблице результатов отдельно считаются damage, effective healing, prevented damage, aggro uptime и poise contribution.")

    section_transition(doc)
    add_heading(doc, "12. Реализация по этапам", 1)
    roadmap = [
        ("Этап 0 — выполнен", "Hunter hotfix и safety pass", "Отказы, multi-hit, ресурсы, пассивы, короткий shield, interruptible casts"),
        ("Этап 1", "CombatProfile + BalanceResolver", "Убрать разбросанные instanceof/магические числа; тесты профилей"),
        ("Этап 2", "Duel/Arena manager", "Snapshot, normalized template, round reset, no-climb arena"),
        ("Этап 3", "Guild War + open-world opt-in", "Compression, subgroup support cap, objectives, anti-camp"),
        ("Этап 4", "Entity tags/data maps + poise", "Совместимость с future mobs/bosses mod"),
        ("Этап 5", "Equipment/jewelry", "PvE main progression; PvP normalization/compression"),
        ("Этап 6", "Priest + migration Heal", "Refund/migrate Mage rank и loadout без потери очков"),
        ("Этап 7", "Summoner", "После стабильного boss API и performance tests"),
    ]
    add_table(doc, ["Порядок", "Результат", "Definition of done"], roadmap,
              [1800, 2800, 4760], font_size=8.8)

    add_heading(doc, "12.1. Миграция Heal из Mage", 2)
    for item in (
        "Сохранить старый ранг HEAL в migration field и вернуть соответствующие skill points Магу.",
        "Удалить HEAL из Mage loadout slots; показать одноразовое сообщение о возврате очков.",
        "После выбора Priest предложить вложить возвращённые очки в Mend/Healing branch, но не делать это автоматически.",
        "До релиза Priest оставить HEAL как legacy, а не удалять ID: иначе старые миры получат сломанные данные.",
    ):
        add_list_item(doc, item, number_id)

    section_transition(doc)
    add_heading(doc, "13. Телеметрия и критерии принятия", 1)
    add_table(
        doc,
        ["Метрика", "Сегментация", "Цель / сигнал проблемы"],
        [
            ("Matchup win rate", "1v1, equal template, mastery bracket", "45–55% после достаточной выборки"),
            ("Median TTK", "Класс × класс × карта", "14–22 с; fast deaths <4 с — редкое полное combo"),
            ("Mobility success", "Chase/escape, stamina spent", "Hunter не top damage и top escape одновременно"),
            ("Cast outcome", "Meteor/Sky: hit, miss, interrupt", "Высокая сила коррелирует с telegraph success"),
            ("CC uptime", "Hard/soft/mobility lock", "Hard CC <15% времени боя"),
            ("Support contribution", "Effective heal, prevent, cleanse", "Stack Priest имеет diminishing value"),
            ("GvG composition", "5v5, 10v10, 20v20", "Ни один класс не >45% оптимального состава"),
            ("PvE boss", "Solo/party/tier", "DPS-class TTK ±15%; utility измеряется отдельно"),
            ("Newcomer safety", "Повторные смерти от одного killer", "Почти 0 вне добровольного War Mode"),
        ],
        [2300, 3000, 4060],
        font_size=8.6,
    )
    add_heading(doc, "13.1. Минимальный тестовый цикл", 2)
    test_number_id = add_numbering(doc, "decimal", "%1.")
    for item in (
        "Не менее 30 дуэлей каждой пары классов на одинаковом template; затем повторить для novice и mastered players.",
        "Отдельные 5v5 и 10v10 сценарии: choke, open field, vertical map, objective defense.",
        "PvE: zombie wave, mixed ranged wave, vanilla boss и один data-driven mock boss с break bar.",
        "После патча менять одну ось за раз: damage, mobility, sustain или control. Иначе нельзя понять причинность.",
        "Публиковать patch note с гипотезой и ожидаемой метрикой, а не только список процентов.",
    ):
        add_list_item(doc, item, test_number_id)

    add_callout(doc, "Когда баланс считается достаточно хорошим", "Игроки выбирают класс из-за роли и удовольствия, а не потому, что один набор одновременно лучше убивает, убегает и переживает ошибку. Небольшие отклонения допустимы; отсутствие контригры — нет.", fill=PALE_BLUE, accent=BLUE)

    section_transition(doc)
    add_heading(doc, "14. Источники и границы переноса", 1)
    add_para(doc, "Использованы первичные или официально поддерживаемые источники. Числа чужих игр служат свидетельством паттернов, а не готовыми настройками Minecraft.")
    sources = [
        ("ArenaNet — разделение PvE/PvP/WvW balance", "https://www.guildwars2.com/en/news/updates-from-the-systems-team/"),
        ("ArenaNet — Defiance bar и telegraph босса", "https://www.guildwars2.com/en/news/meet-the-wyvern-in-guild-wars-2-heart-of-thorns/"),
        ("Guild Wars 2 — structured PvP normalization", "https://wiki.guildwars2.com/wiki/SPvP"),
        ("Blizzard — PvP gear normalization philosophy", "https://worldofwarcraft.blizzard.com/en-us/news/20119625/dev-watercooler-gearing-up-for-legion-pvp"),
        ("Blizzard — War Mode и PvP talents", "https://worldofwarcraft.blizzard.com/en-us/news/21901729"),
        ("Blizzard — crowd-control pruning", "https://worldofwarcraft.blizzard.com/en-us/news/13107743"),
        ("Riot — Champion Balance Framework", "https://www.leagueoflegends.com/en-us/news/dev/dev-champion-balance-framework/"),
        ("Riot — Durability Update", "https://www.leagueoflegends.com/en-us/news/dev/quick-gameplay-thoughts-5-6/"),
        ("Riot — зачем нужны mana/secondary resources", "https://www.leagueoflegends.com/en-us/news/dev/ask-riot-manaless-champions/"),
        ("Valve — Drow Ranger", "https://www.dota2.com/hero/drowranger"),
        ("Valve — Windranger", "https://www.dota2.com/hero/windranger"),
        ("Valve — Sven", "https://www.dota2.com/hero/sven"),
        ("Valve — Ursa", "https://www.dota2.com/hero/ursa"),
        ("Valve — Invoker live data", "https://www.dota2.com/datafeed/herodata?hero_id=74&language=english"),
        ("Valve — Oracle", "https://www.dota2.com/hero/oracle"),
        ("Valve — Chen", "https://www.dota2.com/hero/chen"),
        ("Valve — Lone Druid", "https://www.dota2.com/hero/lonedruid"),
        ("NeoForge — EntityType tags", "https://docs.neoforged.net/docs/resources/server/tags/"),
        ("NeoForge — reloadable data maps", "https://docs.neoforged.net/docs/1.21.5/resources/server/datamaps/"),
    ]
    for idx, (label, url) in enumerate(sources, start=1):
        p = doc.add_paragraph()
        style_para(p, after=4, line=1.15)
        set_run(p.add_run(f"{idx}. "), size=9.5, color=MUTED, bold=True)
        add_hyperlink(p, label, url)

    add_heading(doc, "14.1. Ограничения отчёта", 2)
    add_para(doc, "В проекте пока нет достаточной боевой телеметрии, поэтому предложенные проценты являются alpha-baseline. После появления CombatProfile и тестового harness значения должны перейти в reloadable data files и корректироваться по матчам, а не по впечатлению от одного боя.")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
