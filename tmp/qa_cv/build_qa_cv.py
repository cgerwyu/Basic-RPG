from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path(r"C:\Users\cdald\Desktop\basicrpgclasses-26.2\output\qa_cv")
FONT = "Liberation Sans"
BLACK = RGBColor(0, 0, 0)
GREY = "A6A6A6"


def set_font(run, size, bold=False, color=BLACK):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def set_para(paragraph, before=0, after=0, line=1.0, left=0, keep=False):
    f = paragraph.paragraph_format
    f.space_before = Pt(before)
    f.space_after = Pt(after)
    f.line_spacing = line
    f.left_indent = Pt(left)
    f.keep_with_next = keep


def shade_border_bottom(paragraph):
    ppr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), GREY)
    borders.append(bottom)
    ppr.append(borders)


def remove_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = OxmlElement(f"w:{name}")
        edge.set(qn("w:val"), "nil")
        borders.append(edge)
    tc_pr.append(borders)
    margins = OxmlElement("w:tcMar")
    for side in ("top", "left", "bottom", "right"):
        item = OxmlElement(f"w:{side}")
        item.set(qn("w:w"), "0")
        item.set(qn("w:type"), "dxa")
        margins.append(item)
    tc_pr.append(margins)


def add_text(paragraph, text, size=10.5, bold=False, before=0, after=0, line=1.08, keep=False):
    set_para(paragraph, before, after, line, keep=keep)
    run = paragraph.add_run(text)
    set_font(run, size, bold)
    return paragraph


def add_labeled(paragraph, label, text, size=10.5, before=0, after=0, line=1.08):
    set_para(paragraph, before, after, line)
    first = paragraph.add_run(label)
    set_font(first, size, True)
    rest = paragraph.add_run(text)
    set_font(rest, size, False)


def add_section(doc, title, spacing_before=8):
    p = doc.add_paragraph()
    set_para(p, before=spacing_before, after=3, line=1.0, keep=True)
    run = p.add_run(title)
    set_font(run, 12.5, True)
    shade_border_bottom(p)
    return p


def add_entry(doc, year, title, lines):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.allow_autofit = False
    left, right = table.rows[0].cells
    left.width = Cm(2.45)
    right.width = Cm(14.65)
    for cell in (left, right):
        remove_cell_borders(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    p_left = left.paragraphs[0]
    add_text(p_left, year, 10.5, before=0, after=0, line=1.05)
    p_title = right.paragraphs[0]
    add_text(p_title, title, 10.5, True, before=0, after=1, line=1.05, keep=True)
    for kind, text in lines:
        p = right.add_paragraph()
        if kind == "body":
            add_text(p, text, 10.5, before=0, after=1, line=1.08)
        elif kind == "bullet":
            add_text(p, "- " + text, 10.5, before=0, after=1, line=1.08,)
            p.paragraph_format.left_indent = Pt(4)
        elif kind == "label":
            label, value = text
            add_labeled(p, label, value, 10.5, before=0, after=1, line=1.08)
    # Avoid Word keeping a table row together when it needs to flow normally.
    trpr = table.rows[0]._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    trpr.append(cant_split)


def make_cv(language):
    ru = language == "ru"
    d = Document()
    section = d.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.55)
    section.right_margin = Cm(1.55)
    section.header_distance = Cm(0.5)
    section.footer_distance = Cm(0.5)

    normal = d.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)

    # Header
    p = d.add_paragraph()
    set_para(p, after=6, line=1.0)
    name = p.add_run("Aldar Uliumdzhiev")
    set_font(name, 19, True)

    contacts = [
        ("Tel.: ", "+7 995 551 8149"),
        ("Email: ", "aldar2003@mail.ru"),
        ("GitHub: ", "github.com/cgerwyu"),
        ("LinkedIn: ", "linkedin.com/in/aldar-uliumdzhiev-149a561b8"),
    ]
    for label, value in contacts:
        p = d.add_paragraph()
        set_para(p, after=0, line=1.0)
        a = p.add_run(label)
        set_font(a, 10.5)
        b = p.add_run(value)
        set_font(b, 10.5, color=RGBColor(0, 102, 204) if label != "Tel.: " else BLACK)
        if label != "Tel.: ":
            b.font.underline = True

    p = d.add_paragraph()
    target = "Target position: Junior QA Engineer (Manual)" if not ru else "Целевая позиция: Junior QA Engineer (Manual)"
    add_text(p, target, 10.5, True, before=7, after=6, line=1.0)

    # Profile
    add_section(d, "PROFILE" if not ru else "ПРОФИЛЬ", spacing_before=0)
    profile = (
        [
            "Recent graduate in Applied Informatics (Web Development) with a technical background in Java and web applications.",
            "Seeking an entry-level Manual QA Engineer role to apply analytical thinking, attention to detail and hands-on understanding of client-server products.",
            "Interested in testing user flows, APIs and data-driven features; ready to build practical QA experience in a product team.",
        ] if not ru else [
            "Выпускник направления «Прикладная информатика (веб-разработка)» с техническим бэкграундом в Java и веб-приложениях.",
            "Ищу стартовую позицию Manual QA Engineer, чтобы применить аналитическое мышление, внимательность к деталям и понимание клиент-серверных продуктов.",
            "Интересуюсь тестированием пользовательских сценариев, API и функций, работающих с данными; готов развивать практический QA-опыт в продуктовой команде.",
        ]
    )
    for text in profile:
        p = d.add_paragraph()
        add_text(p, text, 10.5, before=0, after=2, line=1.1)

    # Project experience
    add_section(d, "PET PROJECTS" if not ru else "ПЕТ-ПРОЕКТЫ", spacing_before=6)
    if not ru:
        rpg_title = "Minecraft RPG Mod"
        rpg_lines = [
            ("body", "Personal project in development: an RPG mod for Minecraft built with AI-assisted tools."),
            ("body", "Adds custom RPG mechanics and gameplay content to make player progression more engaging."),
            ("label", ("GitHub: ", "github.com")),
        ]
        project_title = "UFC Fight Parser Telegram Bot"
        project_lines = [
            ("body", "Java Telegram bot that retrieves upcoming UFC events and sends users schedules and reminders."),
            ("label", ("GitHub: ", "github.com/cgerwyu/UFC-Fight-Parser-TG-bot/tree/develop")),
            ("label", ("Stack: ", "Java, Spring Boot, Jsoup, Telegram Bot API, PostgreSQL, Docker Compose")),
        ]
    else:
        rpg_title = "RPG-мод для Minecraft"
        rpg_lines = [
            ("body", "Личный проект в разработке: RPG-мод для Minecraft, создаваемый с использованием AI-инструментов."),
            ("body", "Добавляет пользовательские RPG-механики и игровой контент, делая прогрессию игрока более вовлекающей."),
            ("label", ("GitHub: ", "github.com")),
        ]
        project_title = "Telegram-бот UFC Fight Parser"
        project_lines = [
            ("body", "Java Telegram-бот, который получает данные о ближайших турнирах UFC и отправляет пользователям расписание и напоминания."),
            ("label", ("GitHub: ", "github.com/cgerwyu/UFC-Fight-Parser-TG-bot/tree/develop")),
            ("label", ("Стек: ", "Java, Spring Boot, Jsoup, Telegram Bot API, PostgreSQL, Docker Compose")),
        ]
    add_entry(d, "2026", rpg_title, rpg_lines)
    add_entry(d, "2026", project_title, project_lines)

    # Education
    add_section(d, "EDUCATION" if not ru else "ОБРАЗОВАНИЕ", spacing_before=7)
    edu_title = (
        "Russian Presidential Academy of National Economy and Public Administration (RANEPA)"
        if not ru else "Российская академия народного хозяйства и государственной службы при Президенте РФ (РАНХиГС)"
    )
    edu_line = (
        "Bachelor's degree, Applied Informatics (Web Development), graduated July 2026, average grade: 4.06/5.00 (Russian grading scale)"
        if not ru else "Бакалавриат, прикладная информатика (веб-разработка), окончил в июле 2026, средний балл: 4,06 из 5,00 (российская шкала оценок)"
    )
    add_entry(d, "2022-2026", edu_title, [("body", edu_line)])

    # Course
    course_head = "PROFESSIONAL DEVELOPMENT, COURSES" if not ru else "ПРОФЕССИОНАЛЬНОЕ РАЗВИТИЕ, КУРСЫ"
    add_section(d, course_head, spacing_before=7)
    course_title = "Yandex Practicum - Java Developer: Extended Program" if not ru else "Яндекс Практикум - Java-разработчик: расширенная программа"
    course_lines = (
        [
            ("body", "Professional training course, 855 hours."),
            ("bullet", "Topics covered: Java Core, Spring Boot, databases, ORM, stream processing applications, cloud-native applications and final project."),
        ] if not ru else [
            ("body", "Профессиональная программа обучения, 855 часов."),
            ("bullet", "Темы: Java Core, Spring Boot, базы данных, ORM, приложения потоковой обработки, cloud-native приложения и финальный проект."),
        ]
    )
    add_entry(d, "2025", course_title, course_lines)

    # Skills
    add_section(d, "KEY SKILLS" if not ru else "КЛЮЧЕВЫЕ НАВЫКИ", spacing_before=7)
    if not ru:
        skills = [
            ("Languages: ", "Russian - Native; English - Upper-Intermediate"),
            ("Manual QA: ", "motivated to develop test design, test cases, checklists, bug reporting and regression testing skills"),
            ("Product & technical background: ", "user flows, REST API, client-server architecture, SQL, PostgreSQL, data persistence"),
            ("Tools: ", "Postman, Swagger/OpenAPI, Git, GitHub, Docker, Docker Compose, IntelliJ IDEA"),
            ("Development: ", "Java, OOP, Spring Boot, Spring MVC, Spring Data JPA, JUnit and Mockito basics"),
        ]
    else:
        skills = [
            ("Языки: ", "русский - родной; английский - Upper-Intermediate"),
            ("Ручное QA: ", "мотивирован развивать навыки тест-дизайна, тест-кейсов, чек-листов, баг-репортов и регрессионного тестирования"),
            ("Продуктовая и техническая база: ", "пользовательские сценарии, REST API, клиент-серверная архитектура, SQL, PostgreSQL, сохранение данных"),
            ("Инструменты: ", "Postman, Swagger/OpenAPI, Git, GitHub, Docker, Docker Compose, IntelliJ IDEA"),
            ("Разработка: ", "Java, ООП, Spring Boot, Spring MVC, Spring Data JPA, основы JUnit и Mockito"),
        ]
    for label, text in skills:
        p = d.add_paragraph()
        add_labeled(p, label, text, 10.5, before=0, after=1, line=1.06)

    filename = "Aldar_Uliumdzhiev_Junior_QA_Engineer_Manual_RU.docx" if ru else "Aldar_Uliumdzhiev_Junior_QA_Engineer_Manual_EN.docx"
    d.save(OUT / filename)


if __name__ == "__main__":
    OUT.mkdir(parents=True, exist_ok=True)
    make_cv("en")
    make_cv("ru")
